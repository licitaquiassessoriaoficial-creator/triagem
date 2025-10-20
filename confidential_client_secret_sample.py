import argparse
import hashlib
import io
import json
import re
import sys
import time
import unicodedata
from pathlib import Path

import msal
import requests
from requests.adapters import HTTPAdapter
from urllib3.util.retry import Retry

# Garante que a saída padrão será UTF-8
if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8")


def make_session(max_retries: int) -> requests.Session:
    s = requests.Session()
    retry = Retry(
        total=max_retries,
        connect=max_retries,
        read=max_retries,
        backoff_factor=1.5,
        status_forcelist=[408, 429, 500, 502, 503, 504],
        allowed_methods=frozenset(["GET", "POST"]),
        raise_on_status=False,
        respect_retry_after_header=True,
    )
    adapter = HTTPAdapter(max_retries=retry, pool_maxsize=10)
    s.mount("https://", adapter)
    s.mount("http://", adapter)
    return s


# OCR opcional
try:
    from pdf2image import convert_from_bytes
    import pytesseract
    from PIL import Image
    HAVE_OCR = True
except Exception:
    convert_from_bytes = None
    pytesseract = None
    Image = None
    HAVE_OCR = False

OCR_LANG = "por+eng"
MIN_TEXT_CHARS = 80


def safe_print(msg: str):
    try:
        print(msg, flush=True)
    except UnicodeEncodeError:
        sys.stdout.buffer.write(
            (str(msg) + "\n").encode("utf-8", errors="replace")
        )
        sys.stdout.flush()


_SOFT_HYPHEN = "\u00AD"


def _normalize(s: str) -> str:
    s = s.replace(_SOFT_HYPHEN, "")
    s = re.sub(r'(?<=\w)-\s+(?=\w)', '', s)
    s = s.lower()
    s = ''.join(
        c for c in unicodedata.normalize('NFD', s)
        if unicodedata.category(c) != 'Mn'
    )
    s = re.sub(r'\s+', ' ', s).strip()
    return s


def normalize_text(s):
    s = s.lower()
    s = ''.join(
        c for c in unicodedata.normalize('NFD', s)
        if unicodedata.category(c) != 'Mn'
    )
    s = re.sub(r'\s+', ' ', s).strip()
    return s


def _has_exact_phrase(texto: str, frase: str) -> bool:
    t = _normalize(texto)
    p = _normalize(frase)
    return re.search(r'\b' + re.escape(p) + r'\b', t) is not None


def load_config(path: str) -> dict:
    with open(path, "r", encoding="utf-8") as f:
        return json.load(f)


def get_token(cfg: dict) -> str:
    app = msal.ConfidentialClientApplication(
        client_id=cfg["client_id"],
        authority=cfg["authority"],
        client_credential=cfg["secret"],
    )
    scopes = cfg.get("scope", ["https://graph.microsoft.com/.default"])
    if isinstance(scopes, str):
        scopes = [scopes]
    result = app.acquire_token_silent(scopes=scopes, account=None)
    if not result:
        safe_print("[INFO] Nenhum token no cache. Solicitando novo...")
        result = app.acquire_token_for_client(scopes=scopes)
    if "access_token" not in result:
        error_msg = (
            f"[ERRO] Falha ao obter token: {result.get('error')}: "
            f"{result.get('error_description')}"
        )
        safe_print(error_msg)
        sys.exit(2)
    return result["access_token"]


def fetch_messages(
    endpoint: str,
    token: str,
    timeout: int = 60,
    max_retries: int = 5
):
    sess = make_session(max_retries)
    headers = {
        "Authorization": f"Bearer {token}",
        "Accept": "application/json"
    }
    url = endpoint
    items = []
    page = 1
    while url:
        attempt = 0
        while True:
            attempt += 1
            try:
                safe_print(f"Buscando página {page}: {url}")
                resp = sess.get(url, headers=headers, timeout=(10, timeout))
                if resp.status_code == 429:
                    ra = int(resp.headers.get("Retry-After", "5"))
                    safe_print(
                        f"[RETRY] 429 Too Many Requests – aguardando {ra}s"
                    )
                    time.sleep(ra)
                    continue
                if not resp.ok:
                    msg = f"[WARN] Graph retornou {resp.status_code}: "
                    msg += f"{resp.text[:300]}"
                    safe_print(msg)
                    resp.raise_for_status()
                data = resp.json()
                batch = data.get("value", [])
                items.extend(batch)
                url = data.get("@odata.nextLink")
                msg = f"Página {page} recebida, msgs={len(batch)}, "
                msg += f"acumulado={len(items)}"
                safe_print(msg)
                page += 1
                break
            except (
                requests.exceptions.ReadTimeout,
                requests.exceptions.ConnectTimeout,
                requests.exceptions.ConnectionError
            ) as e:
                if attempt <= max_retries:
                    wait = min(2 ** attempt, 30)
                    msg = f"[RETRY] {type(e).__name__}: tentando de novo em "
                    msg += f"{wait}s (tentativa {attempt}/{max_retries})"
                    safe_print(msg)
                    time.sleep(wait)
                    continue
                else:
                    msg = f"[ERRO] Falhou após {max_retries} tentativas: {e}"
                    safe_print(msg)
                    raise
    return items


def list_attachments(user_email, msg_id, token):
    base_url = "https://graph.microsoft.com/v1.0/users"
    url = f"{base_url}/{user_email}/messages/{msg_id}/attachments"
    headers = {"Authorization": f"Bearer {token}"}
    resp = requests.get(url, headers=headers)
    if resp.status_code != 200:
        safe_print(f"[WARN] Falha ao listar anexos: {resp.text}")
        return []
    return resp.json().get("value", [])


def download_attachment(user_email, msg_id, att_id, token):
    base_url = "https://graph.microsoft.com/v1.0/users"
    url = f"{base_url}/{user_email}/messages/{msg_id}/attachments/{att_id}"
    headers = {"Authorization": f"Bearer {token}"}
    resp = requests.get(url, headers=headers)
    if resp.status_code != 200:
        safe_print(f"[WARN] Falha ao baixar anexo {att_id}: {resp.text}")
        return None, None, None
    att = resp.json()
    if att.get('@odata.type') == '#microsoft.graph.fileAttachment':
        fname = att.get('name', att_id)
        ctype = att.get('contentType', '')
        import base64
        content = att['contentBytes']
        if isinstance(content, str):
            try:
                data = base64.b64decode(content)
            except Exception:
                data = content.encode('utf-8')
        else:
            data = content
        return fname, data, ctype
    return None, None, None


_ILLEGAL = r'[\\/:*?"<>|]'


def safe_name(name: str, fallback_ext: str = "") -> str:
    name = (name or "anexo").strip()
    base, dot, ext = name.rpartition('.')
    if not dot:
        base, ext = name, fallback_ext.lstrip('.')
    base = re.sub(_ILLEGAL, "_", base) or "arquivo"
    ext = re.sub(_ILLEGAL, "", ext.lower())
    full = f"{base}.{ext}" if ext else base
    if len(full) > 120:
        h = hashlib.sha1(full.encode('utf-8')).hexdigest()[:8]
        keep = 120 - (len(ext) + (1 if ext else 0) + 9)
        base = base[:max(10, keep)]
        full = f"{base}_{h}.{ext}" if ext else f"{base}_{h}"
    return full


def save_bytes(tmp_dir: Path, filename: str, data: bytes) -> Path:
    tmp_dir.mkdir(parents=True, exist_ok=True)
    path = tmp_dir / safe_name(filename)
    with open(path, "wb") as f:
        f.write(data)
    return path


UNSUPPORTED_EXT = (
    ".zip", ".rar", ".7z", ".msg", ".eml", ".xls", ".xlsx", ".ppt", ".pptx"
)


def extract_text_any(fname: str, ctype: str, data: bytes) -> tuple[str, bool]:
    # PDF
    if fname.lower().endswith(".pdf") or "pdf" in (ctype or "").lower():
        text = ""
        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(io.BytesIO(data))
            for p in reader.pages:
                text += (p.extract_text() or "") + "\n"
        except Exception:
            text = ""
        if len(text.strip()) >= MIN_TEXT_CHARS:
            return text, False
        # fallback OCR
        if HAVE_OCR and convert_from_bytes is not None:
            try:
                pages = convert_from_bytes(data, dpi=300)
                ocr_text = []
                for img in pages:
                    ocr_result = pytesseract.image_to_string(
                        img, lang=OCR_LANG
                    )
                    ocr_text.append(ocr_result)
                return "\n".join(ocr_text), True
            except Exception:
                pass
        return text, False

    # DOCX
    is_docx = (fname.lower().endswith(".docx") or
               "officedocument.wordprocessingml.document" in
               (ctype or "").lower())
    if is_docx:
        try:
            import tempfile
            import docx
            with tempfile.NamedTemporaryFile(
                suffix=".docx", delete=True
            ) as tmp:
                tmp.write(data)
                tmp.flush()
                d = docx.Document(tmp.name)
                return "\n".join(p.text for p in d.paragraphs), False
        except Exception:
            return "", False

    # TXT
    if fname.lower().endswith(".txt") or "text/plain" in (ctype or "").lower():
        try:
            return data.decode("utf-8", errors="ignore"), False
        except Exception:
            return "", False

    # Imagens (PNG/JPG/TIF)
    image_exts = (".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp")
    is_image = (any(fname.lower().endswith(ext) for ext in image_exts) or
                (ctype or "").lower().startswith("image/"))
    if is_image:
        if HAVE_OCR and Image is not None and pytesseract is not None:
            try:
                img = Image.open(io.BytesIO(data))
                return pytesseract.image_to_string(img, lang=OCR_LANG), True
            except Exception:
                return "", False
        else:
            return "", False

    return "", False


def extract_text_from_pdf(path):
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(path)
        return "\n".join((p.extract_text() or "") for p in reader.pages)
    except Exception:
        print("[WARN] Falha ao extrair texto de", path)
        return ""


def extract_text_from_docx(path):
    try:
        import docx
        d = docx.Document(path)
        return "\n".join(p.text for p in d.paragraphs)
    except Exception:
        print("[WARN] Falha ao extrair texto de", path)
        return ""


def get_resume_text(path):
    path_l = path.lower()
    if path_l.endswith(".pdf"):
        return extract_text_from_pdf(path)
    elif path_l.endswith(".docx"):
        return extract_text_from_docx(path)
    elif path_l.endswith(".txt"):
        try:
            with open(path, "r", encoding="utf-8") as f:
                return f.read()
        except Exception:
            print("[WARN] Falha ao extrair texto de", path)
            return ""
    else:
        print("[WARN] Extensão não suportada para", path)
        return ""


FORMAÇÃO_SYNONYMS = {
    "farmacia": [
        "farmacia", "farmácia", "farmaceutico", "farmacêutico", "pharmacy"
    ],
    "biomedicina": ["biomedicina", "biomédico", "biomédica"],
    "quimica": [
        "quimica", "química", "quimico", "química industrial", "chemistry"
    ],
}


def expand_formacoes(user_terms):
    result = set()
    for term in user_terms:
        norm = normalize_text(term)
        result.add(norm)
        for key, syns in FORMAÇÃO_SYNONYMS.items():
            if norm == key or norm in syns:
                result.update([normalize_text(s) for s in syns])
    return result


def has_formacao_in_text(cv_text, formacoes):
    cv_norm = normalize_text(cv_text)
    expanded = expand_formacoes(formacoes)
    for term in expanded:
        if term in cv_norm:
            return True
    return False


def candidato_aprovado(texto_cv, palavras, formacoes):
    texto_norm = normalize_text(texto_cv)
    formacoes_encontradas = set()
    if palavras:
        if not any(normalize_text(p) in texto_norm for p in palavras):
            return False, formacoes_encontradas
    if formacoes:
        expanded = expand_formacoes(formacoes)
        for term in expanded:
            if term in texto_norm:
                formacoes_encontradas.add(term)
        if not formacoes_encontradas:
            return False, formacoes_encontradas
    return True, formacoes_encontradas


def teste_formacao():
    cvs = [
        "Graduado em farmácia pela USP, experiência em química",
        "Formação em biomedicina, trabalho em laboratório",
        "Engenheiro civil, sem experiência em área médica",
    ]
    formacoes = ["farmácia", "biomedicina"]
    for i, cv in enumerate(cvs, 1):
        encontrados = set()
        if formacoes:
            expanded = expand_formacoes(formacoes)
            cv_norm = normalize_text(cv)
            for termo in expanded:
                if termo in cv_norm:
                    encontrados.add(termo)
        aprovado = bool(encontrados)
        msg = (
            f"CV {i}: {'APROVADO' if aprovado else 'REPROVADO'} | "
            f"Encontradas: "
            f"{', '.join(encontrados) if encontrados else 'nenhuma'}"
        )
        print(msg)


def main():
    parser = argparse.ArgumentParser(
        description="Triagem de currículos via Microsoft Graph"
    )
    parser.add_argument("vaga_desc", help="Descrição da vaga")
    parser.add_argument(
        "keywords", help="Palavras-chave separadas por vírgula"
    )
    parser.add_argument("--negativas", help="Palavras negativas")
    parser.add_argument(
        "--http-timeout",
        type=int,
        default=60,
        help="Timeout de leitura (segundos)"
    )
    parser.add_argument(
        "--max-retries",
        type=int,
        default=5,
        help="Máximo de tentativas em falha de rede"
    )
    parser.add_argument(
        "--formacoes",
        default="",
        help="Lista de formações separadas por vírgula"
    )
    parser.add_argument(
        "--teste-formacao",
        action="store_true",
        help="Executa teste de formação em memória"
    )

    args = parser.parse_args()

    if args.teste_formacao:
        # Executar teste inline
        cvs = [
            "Graduado em farmácia pela USP, experiência em química",
            "Formação em biomedicina, trabalho em laboratório",
            "Engenheiro civil, sem experiência em área médica",
        ]
        formacoes = ["farmácia", "biomedicina"]
        for i, cv in enumerate(cvs, 1):
            encontrados = set()
            if formacoes:
                expanded = expand_formacoes(formacoes)
                cv_norm = normalize_text(cv)
                for termo in expanded:
                    if termo in cv_norm:
                        encontrados.add(termo)
            aprovado = bool(encontrados)
            msg = (
                f"CV {i}: {'APROVADO' if aprovado else 'REPROVADO'} | "
                f"Encontradas: "
                f"{', '.join(encontrados) if encontrados else 'nenhuma'}"
            )
            print(msg)
        return

    # Implementação principal
    params = load_config("parameters.json")

    # Processamento das formações
    cli_form = [
        s.strip() for s in (args.formacoes or "").split(",") if s.strip()
    ]
    json_form = [
        s.strip() for s in (params.get("formacoes", "") or "").split(",")
        if s.strip()
    ]
    formacoes = expand_formacoes(cli_form or json_form)

    msg_formacao = (
        f"[INFO] Filtro de formação: "
        f"{', '.join(formacoes) if formacoes else 'vazio'}"
    )
    safe_print(msg_formacao)

    token = get_token(params)
    user_email = params["user_email"]
    endpoint = params["endpoint"]
    base_dir = Path(params.get("base_dir", "aprovados"))
    qual_dir = base_dir / "qualidade"
    tmp_dir = base_dir / "__tmp"

    qual_dir.mkdir(parents=True, exist_ok=True)
    tmp_dir.mkdir(parents=True, exist_ok=True)

    positivas = [args.vaga_desc.strip()] + [
        k.strip() for k in args.keywords.split(",") if k.strip()
    ]
    negativas = (
        [k.strip() for k in args.negativas.split(",") if k.strip()]
        if args.negativas else []
    )

    if not HAVE_OCR:
        msg_ocr = (
            "[WARN] OCR indisponível (instale Tesseract + pdf2image + "
            "Pillow). PDFs/Imagens podem ficar sem texto."
        )
        safe_print(msg_ocr)

    messages = fetch_messages(
        endpoint, token,
        timeout=args.http_timeout,
        max_retries=args.max_retries
    )

    safe_print(f"[INFO] {len(messages)} mensagens obtidas")

    aprovados = []
    for msg in messages:
        try:
            subj = msg.get("subject", "(sem assunto)")
            msg_id = msg["id"]
            msg_from = msg.get("from", {}).get(
                "emailAddress", {}
            ).get("address", "(desconhecido)")
            safe_print(f"[MSG] De: {msg_from} | Assunto: {subj[:50]}")

            atts = list_attachments(user_email, msg_id, token)
            for att in atts:
                try:
                    att_name = att["name"]
                    if any(att_name.lower().endswith(ext)
                           for ext in UNSUPPORTED_EXT):
                        safe_print(
                            f"[SKIP] Extensão não suportada: {att_name}"
                        )
                        continue

                    fname, data, ctype = download_attachment(
                        user_email, msg_id, att['id'], token
                    )
                    if not data:
                        continue

                    local = save_bytes(tmp_dir, fname, data)
                    text, ocr_used = extract_text_any(fname, ctype, data)

                    aprovado, formacoes_encontradas = candidato_aprovado(
                        text, positivas, formacoes
                    )

                    pos_hit = any(
                        normalize_text(k) in normalize_text(text)
                        for k in positivas
                    )
                    neg_hit = any(
                        normalize_text(k) in normalize_text(text)
                        for k in negativas
                    )

                    msg_scan = (
                        f"[SCAN] {local.name} chars={len(text)} "
                        f"ocr={ocr_used} pos_hit={pos_hit} neg_hit={neg_hit}"
                    )
                    safe_print(msg_scan)

                    if aprovado and not neg_hit:
                        qual_path = qual_dir / local.name
                        qual_path.write_bytes(data)
                        formacoes_str = (
                            ", ".join(formacoes_encontradas)
                            if formacoes_encontradas else ""
                        )
                        aprovados.append({
                            "nome": fname,
                            "from": msg_from,
                            "subject": subj,
                            "formacoes_encontradas": formacoes_str,
                            "aprovado": True
                        })
                        safe_print(f"[OK] {fname} salvo em qualidade/")
                    else:
                        safe_print(f"[SKIP] {fname} - reprovado/negativas")
                except Exception as e:
                    safe_print(f"[ERRO] Anexo {att['name']}: {e}")
        except Exception as e:
            safe_print(f"[ERRO] Mensagem {msg.get('id', '?')}: {e}")

    # Salvar resultados
    import csv
    csv_path = base_dir / "aprovados.csv"
    json_path = base_dir / "aprovados.json"

    with open(csv_path, "w", newline="", encoding="utf-8") as csvfile:
        if aprovados:
            writer = csv.DictWriter(csvfile, fieldnames=aprovados[0].keys())
            writer.writeheader()
            writer.writerows(aprovados)

    with open(json_path, "w", encoding="utf-8") as jsonfile:
        json.dump(aprovados, jsonfile, indent=2, ensure_ascii=False)

    safe_print(f"[INFO] {len(aprovados)} candidatos aprovados")
    safe_print(f"CSV salvo em: {csv_path}")
    safe_print(f"JSON salvo em: {json_path}")


if __name__ == "__main__":
    main()
