"""
Sistema de triagem de currículos via Microsoft Graph API.
Processa apenas emails não lidos e marca como lidos após processamento.
"""

import argparse
import collections
import csv  # Usado para csv.DictWriter na linha 650
import hashlib
import io
import json
import re
import sys
import time
import unicodedata
from pathlib import Path

import msal
import requests
from requests.adapters import HTTPAdapter
from urllib3.util.retry import Retry

# Garante que a saída padrão será UTF-8
if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8")

# OCR opcional
try:
    from pdf2image import convert_from_bytes
    import pytesseract
    from PIL import Image
    HAVE_OCR = True
except Exception:
    convert_from_bytes = None
    pytesseract = None
    Image = None
    HAVE_OCR = False

OCR_LANG = "por+eng"
MIN_TEXT_CHARS = 80

UNSUPPORTED_EXT = (
    ".zip", ".rar", ".7z", ".msg", ".eml",
    ".xls", ".xlsx", ".ppt", ".pptx"
)

FORMACAO_SYNONYMS = {
    "farmacia": [
        "farmacia", "farmácia", "farmaceutico", "farmacêutico", "pharmacy"
    ],
    "biomedicina": ["biomedicina", "biomédico", "biomédica"],
    "quimica": [
        "quimica", "química", "quimico", "química industrial", "chemistry"
    ],
}

def make_session(max_retries: int) -> requests.Session:
    """Cria sessão HTTP com retry automático."""
    s = requests.Session()
    retry = Retry(
        total=max_retries,
        connect=max_retries,
        read=max_retries,
        backoff_factor=1.5,
        status_forcelist=[408, 429, 500, 502, 503, 504],
        allowed_methods=frozenset(["GET", "POST"]),
        raise_on_status=False,
        respect_retry_after_header=True,
    )
    adapter = HTTPAdapter(max_retries=retry, pool_maxsize=10)
    s.mount("https://", adapter)
    s.mount("http://", adapter)
    return s

def safe_print(msg: str):
    """Imprime mensagem com encoding seguro."""
    try:
        print(msg, flush=True)
    except UnicodeEncodeError:
        buffer_msg = (str(msg) + "\n").encode("utf-8", errors="replace")
        sys.stdout.buffer.write(buffer_msg)
        sys.stdout.flush()

def normalize_text(s):
    """Normaliza texto removendo acentos e padronizando."""
    s = s.lower()
    s = ''.join(
        c for c in unicodedata.normalize('NFD', s)
        if unicodedata.category(c) != 'Mn'
    )
    s = re.sub(r'\s+', ' ', s).strip()
    return s

def _normalize(s: str) -> str:
    """Normalização avançada de texto."""
    soft_hyphen = "\u00AD"
    s = s.replace(soft_hyphen, "")
    s = re.sub(r'(?<=\w)-\s+(?=\w)', '', s)
    s = s.lower()
    s = ''.join(
        c for c in unicodedata.normalize('NFD', s)
        if unicodedata.category(c) != 'Mn'
    )
    s = re.sub(r'\s+', ' ', s).strip()
    return s

def _has_exact_phrase(texto: str, frase: str) -> bool:
    """Verifica se frase exata existe no texto."""
    t = _normalize(texto)
    p = _normalize(frase)
    return re.search(r'\b' + re.escape(p) + r'\b', t) is not None

def load_config(path: str) -> dict:
    """Carrega configuração do arquivo JSON."""
    with open(path, "r", encoding="utf-8") as f:
        return json.load(f)

def get_token(cfg: dict) -> str:
    """Obtém token de acesso da Microsoft."""
    app = msal.ConfidentialClientApplication(
        client_id=cfg["client_id"],
        authority=cfg["authority"],
        client_credential=cfg["secret"],
    )
    scopes = cfg.get("scope", ["https://graph.microsoft.com/.default"])
    if isinstance(scopes, str):
        scopes = [scopes]

    result = app.acquire_token_silent(scopes=scopes, account=None)
    if not result:
        safe_print("[INFO] Nenhum token no cache. Solicitando novo...")
        result = app.acquire_token_for_client(scopes=scopes)

    if "access_token" not in result:
        error_msg = (
            f"[ERRO] Falha ao obter token: {result.get('error')}: "
            f"{result.get('error_description')}"
        )
        safe_print(error_msg)
        sys.exit(2)
    return result["access_token"]

def fetch_messages(endpoint: str, token: str, timeout: int = 60,
                   max_retries: int = 5):
    """Busca mensagens do Microsoft Graph API."""
    sess = make_session(max_retries)
    headers = {
        "Authorization": f"Bearer {token}",
        "Accept": "application/json"
    }
    url = endpoint
    items = []
    page = 1

    while url:
        attempt = 0
        while True:
            attempt += 1
            try:
                safe_print(f"Buscando página {page}: {url}")
                resp = sess.get(url, headers=headers, timeout=(10, timeout))

                if resp.status_code == 429:
                    ra = int(resp.headers.get("Retry-After", "5"))
                    safe_print(f"[RETRY] 429 - aguardando {ra}s")
                    time.sleep(ra)
                    continue

                if not resp.ok:
                    warn_msg = (
                        f"[WARN] Graph retornou {resp.status_code}: "
                        f"{resp.text[:300]}"
                    )
                    safe_print(warn_msg)
                    resp.raise_for_status()

                data = resp.json()
                batch = data.get("value", [])
                items.extend(batch)
                url = data.get("@odata.nextLink")

                page_msg = (
                    f"Página {page} recebida, msgs={len(batch)}, "
                    f"acumulado={len(items)}"
                )
                safe_print(page_msg)
                page += 1
                break

            except (
                requests.exceptions.ReadTimeout,
                requests.exceptions.ConnectTimeout,
                requests.exceptions.ConnectionError
            ) as e:
                if attempt <= max_retries:
                    wait = min(2 ** attempt, 30)
                    retry_msg = (
                        f"[RETRY] {type(e).__name__}: tentando novamente em "
                        f"{wait}s (tentativa {attempt}/{max_retries})"
                    )
                    safe_print(retry_msg)
                    time.sleep(wait)
                    continue
                else:
                    error_msg = (
                        f"[ERRO] Falhou após {max_retries} tentativas: {e}"
                    )
                    safe_print(error_msg)
                    raise
    return items

def list_attachments(user_email, msg_id, token):
    """Lista anexos de uma mensagem."""
    base_url = "https://graph.microsoft.com/v1.0"
    url = f"{base_url}/users/{user_email}/messages/{msg_id}/attachments"
    headers = {"Authorization": f"Bearer {token}"}
    resp = requests.get(url, headers=headers)
    if resp.status_code != 200:
        safe_print(f"[WARN] Falha ao listar anexos: {resp.text}")
        return []
    return resp.json().get("value", [])

def mark_email_as_read(user_email, msg_id, token):
    """Marca um email como lido."""
    base_url = "https://graph.microsoft.com/v1.0"
    url = f"{base_url}/users/{user_email}/messages/{msg_id}"
    headers = {
        "Authorization": f"Bearer {token}",
        "Content-Type": "application/json"
    }
    data = {"isRead": True}
    try:
        resp = requests.patch(url, headers=headers, json=data)
        if resp.status_code == 200:
            return True
        else:
            warn_msg = (
                f"[WARN] Falha ao marcar como lido: {resp.status_code}"
            )
            safe_print(warn_msg)
            return False
    except Exception as e:
        safe_print(f"[WARN] Erro ao marcar como lido: {e}")
        return False

def download_attachment(user_email, msg_id, att_id, token):
    """Baixa um anexo específico."""
    base_url = "https://graph.microsoft.com/v1.0"
    url = (
        f"{base_url}/users/{user_email}/messages/"
        f"{msg_id}/attachments/{att_id}"
    )
    headers = {"Authorization": f"Bearer {token}"}
    resp = requests.get(url, headers=headers)

    if resp.status_code != 200:
        safe_print(f"[WARN] Falha ao baixar anexo {att_id}: {resp.text}")
        return None, None, None

    att = resp.json()
    if att.get('@odata.type') == '#microsoft.graph.fileAttachment':
        fname = att.get('name', att_id)
        ctype = att.get('contentType', '')
        import base64
        content = att['contentBytes']
        if isinstance(content, str):
            try:
                data = base64.b64decode(content)
            except Exception:
                data = content.encode('utf-8')
        else:
            data = content
        return fname, data, ctype
    return None, None, None

def safe_name(name: str, fallback_ext: str = "") -> str:
    """Cria nome de arquivo seguro."""
    illegal = r'[\\/:*?"<>|]'
    name = (name or "anexo").strip()
    base, dot, ext = name.rpartition('.')
    if not dot:
        base, ext = name, fallback_ext.lstrip('.')
    base = re.sub(illegal, "_", base) or "arquivo"
    ext = re.sub(illegal, "", ext.lower())
    full = f"{base}.{ext}" if ext else base
    if len(full) > 120:
        h = hashlib.sha1(full.encode('utf-8')).hexdigest()[:8]
        keep = 120 - (len(ext) + (1 if ext else 0) + 9)
        base = base[:max(10, keep)]
        full = f"{base}_{h}.{ext}" if ext else f"{base}_{h}"
    return full

def save_bytes(tmp_dir: Path, filename: str, data: bytes) -> Path:
    """Salva bytes em arquivo."""
    tmp_dir.mkdir(parents=True, exist_ok=True)
    path = tmp_dir / safe_name(filename)
    with open(path, "wb") as f:
        f.write(data)
    return path

def extract_text_any(fname: str, ctype: str, data: bytes) -> tuple:
    """Extrai texto de qualquer tipo de arquivo."""
    # PDF
    if fname.lower().endswith(".pdf") or "pdf" in (ctype or "").lower():
        text = ""
        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(io.BytesIO(data))
            for p in reader.pages:
                text += (p.extract_text() or "") + "\n"
        except Exception:
            text = ""

        if len(text.strip()) >= MIN_TEXT_CHARS:
            return text, False

        # Fallback OCR
        if HAVE_OCR and convert_from_bytes is not None:
            try:
                pages = convert_from_bytes(data, dpi=300)
                ocr_text = []
                for img in pages:
                    ocr_text.append(
                        pytesseract.image_to_string(img, lang=OCR_LANG)
                    )
                return "\n".join(ocr_text), True
            except Exception:
                pass
        return text, False

    # DOCX
    if (fname.lower().endswith(".docx") or
            "officedocument.wordprocessingml.document" in
            (ctype or "").lower()):
        try:
            import tempfile
            import docx
            with tempfile.NamedTemporaryFile(
                suffix=".docx", delete=True
            ) as tmp:
                tmp.write(data)
                tmp.flush()
                d = docx.Document(tmp.name)
                return "\n".join(p.text for p in d.paragraphs), False
        except Exception:
            return "", False

    # TXT
    if (fname.lower().endswith(".txt") or
            "text/plain" in (ctype or "").lower()):
        try:
            return data.decode("utf-8", errors="ignore"), False
        except Exception:
            return "", False

    # Imagens
    img_extensions = (".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp")
    if (any(fname.lower().endswith(ext) for ext in img_extensions) or
            (ctype or "").lower().startswith("image/")):
        if HAVE_OCR and Image is not None and pytesseract is not None:
            try:
                img = Image.open(io.BytesIO(data))
                return pytesseract.image_to_string(img, lang=OCR_LANG), True
            except Exception:
                return "", False
        else:
            return "", False

    return "", False

def get_resume_text(path):
    """Extrai texto de arquivo de currículo."""
    path_l = path.lower()
    if path_l.endswith(".pdf"):
        return extract_text_from_pdf(path)
    elif path_l.endswith(".docx"):
        return extract_text_from_docx(path)
    elif path_l.endswith(".txt"):
        try:
            with open(path, "r", encoding="utf-8") as f:
                return f.read()
        except Exception:
            safe_print(f"[WARN] Falha ao extrair texto de {path}")
            return ""
    else:
        safe_print(f"[WARN] Extensão não suportada para {path}")
        return ""

def extract_text_from_pdf(path):
    """Extrai texto de PDF."""
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(path)
        return "\n".join((p.extract_text() or "") for p in reader.pages)
    except Exception:
        safe_print(f"[WARN] Falha ao extrair texto de {path}")
        return ""

def extract_text_from_docx(path):
    """Extrai texto de DOCX."""
    try:
        import docx
        d = docx.Document(path)
        return "\n".join(p.text for p in d.paragraphs)
    except Exception:
        safe_print(f"[WARN] Falha ao extrair texto de {path}")
        return ""

def expand_formacoes(user_terms):
    """Expande termos de formação com sinônimos."""
    result = set()
    for term in user_terms:
        norm = normalize_text(term)
        result.add(norm)
        for key, syns in FORMACAO_SYNONYMS.items():
            if norm == key or norm in syns:
                result.update([normalize_text(s) for s in syns])
    return result

def has_formacao_in_text(cv_text, formacoes):
    """Verifica se formação está presente no texto."""
    cv_norm = normalize_text(cv_text)
    expanded = expand_formacoes(formacoes)
    for term in expanded:
        if term in cv_norm:
            return True
    return False

def candidato_aprovado(texto_cv, palavras, formacoes):
    """Determina se candidato deve ser aprovado."""
    texto_norm = normalize_text(texto_cv)
    formacoes_encontradas = set()

    if palavras:
        if not any(normalize_text(p) in texto_norm for p in palavras):
            return False, formacoes_encontradas

    if formacoes:
        expanded = expand_formacoes(formacoes)
        for term in expanded:
            if term in texto_norm:
                formacoes_encontradas.add(term)
        if not formacoes_encontradas:
            return False, formacoes_encontradas

    return True, formacoes_encontradas

def teste_formacao():
    """Teste das funcionalidades de formação."""
    textos = [
        "Bacharel em Farmácia pela UFMG...",
        "Graduado em Química Industrial...",
        "Experiência em laboratório; cursando Biomedicina"
    ]
    sets = [
        ["farmácia", "biomedicina"],
        ["química"]
    ]
    for idx, formacoes in enumerate(sets, 1):
        safe_print(f"\n[TESTE] Formações: {formacoes}")
        for i, texto in enumerate(textos, 1):
            aprovado, encontrados = candidato_aprovado(texto, [], formacoes)
            status = 'APROVADO' if aprovado else 'REPROVADO'
            encontradas_str = (
                ', '.join(encontrados) if encontrados else 'nenhuma'
            )
            safe_print(f"CV {i}: {status} | Encontradas: {encontradas_str}")

def main():
    """Função principal."""
    parser = argparse.ArgumentParser()
    parser.add_argument("parameters_json")
    parser.add_argument("vaga_desc")
    parser.add_argument("keywords")
    parser.add_argument("negativas", nargs="?", default="")
    parser.add_argument(
        "--http-timeout", type=int, default=60,
        help="Timeout de leitura (segundos)"
    )
    parser.add_argument(
        "--max-retries", type=int, default=5,
        help="Máximo de tentativas em falha de rede"
    )
    parser.add_argument(
        "--formacoes", default="",
        help="Lista de formações separadas por vírgula"
    )
    parser.add_argument(
        "--teste-formacao", action="store_true",
        help="Executa teste de formação em memória"
    )
    args = parser.parse_args()

    if args.teste_formacao:
        teste_formacao()
        sys.exit(0)

    safe_print("[START] triagem")
    params = load_config(args.parameters_json)

    cli_form = [
        s.strip() for s in (args.formacoes or "").split(",") if s.strip()
    ]
    json_form = [
        s.strip() for s in (params.get("formacoes", "") or "").split(",")
        if s.strip()
    ]
    formacoes = list(dict.fromkeys(cli_form or json_form))

    formacoes_str = ', '.join(formacoes) if formacoes else 'vazio'
    safe_print(f"[INFO] Filtro de formação: {formacoes_str}")

    token = get_token(params)
    endpoint = params["endpoint"]

    m = re.search(r"/users/([^/]+)/messages", endpoint)
    user_email = m.group(1) if m else None
    if not user_email:
        safe_print("[ERRO] Não foi possível extrair user_email do endpoint")
        sys.exit(1)

    vaga_slug = _normalize(args.vaga_desc).replace(' ', '_') or 'vaga'
    base_dir = Path.cwd() / "aprovados" / vaga_slug
    tmp_dir = base_dir / "__tmp"
    base_dir.mkdir(parents=True, exist_ok=True)
    tmp_dir.mkdir(parents=True, exist_ok=True)

    positivas = [args.vaga_desc.strip()] + [
        k.strip() for k in args.keywords.split(",") if k.strip()
    ]
    negativas = [
        k.strip() for k in args.negativas.split(",") if k.strip()
    ] if args.negativas else []

    if not HAVE_OCR:
        msg = "[WARN] OCR indisponível. PDFs/Imagens podem ficar sem texto."
        safe_print(msg)

    # Verificar se está filtrando apenas emails não lidos
    filtro_nao_lidos = "isRead eq false" in endpoint
    if filtro_nao_lidos:
        safe_print("[INFO] Filtrando apenas emails NÃO LIDOS")
    else:
        safe_print("[INFO] Buscando TODOS os emails")

    messages = fetch_messages(
        endpoint, token, timeout=args.http_timeout,
        max_retries=args.max_retries
    )
    safe_print(f"[INFO] emails recebidos: {len(messages)}")

    total_anexos = 0
    total_aprovados = 0
    aprovados_info = []
    formacao_counter = collections.Counter()

    for msg in messages:
        msg_id = msg.get('id')
        if not msg.get('hasAttachments'):
            continue

        atts = list_attachments(user_email, msg_id, token)

        for att in atts:
            name_l = (att.get("name") or "").lower()
            if any(name_l.endswith(ext) for ext in UNSUPPORTED_EXT):
                safe_print(f"[SKIP] {att.get('name')} tipo não suportado")
                continue

            if att.get('@odata.type') != '#microsoft.graph.fileAttachment':
                continue

            fname, data, ctype = download_attachment(
                user_email, msg_id, att['id'], token
            )
            if not fname or not data:
                continue

            total_anexos += 1
            local = save_bytes(tmp_dir, fname, data)
            texto_cv = get_resume_text(str(local))

            aprovado, formacoes_encontradas = candidato_aprovado(
                texto_cv, positivas, formacoes
            )

            if not aprovado and formacoes:
                debug_msg = (
                    f"[DEBUG] Reprovado (formação não encontrada) → {fname}"
                )
                safe_print(debug_msg)

            text, ocr_used = extract_text_any(local.name, ctype, data)
            pos_hit = any(_has_exact_phrase(text, p) for p in positivas)
            neg_hit = any(_has_exact_phrase(text, n) for n in negativas)

            scan_msg = (
                f"[SCAN] {local.name} chars={len(text)} ocr={ocr_used} "
                f"pos_hit={pos_hit} neg_hit={neg_hit}"
            )
            safe_print(scan_msg)

            if aprovado and pos_hit and not neg_hit:
                target = base_dir / local.name
                local.replace(target)
                total_aprovados += 1
                safe_print(f"[MOVE] APPROVED -> {target.name}")

                formacoes_str_encontradas = (
                    ", ".join(formacoes_encontradas)
                    if formacoes_encontradas else ""
                )
                aprovados_info.append({
                    "arquivo": target.name,
                    "formacoes_encontradas": formacoes_str_encontradas
                })
                formacao_counter.update(formacoes_encontradas)
            else:
                local.unlink(missing_ok=True)

        # Marcar email como lido após processar todos os anexos
        if mark_email_as_read(user_email, msg_id, token):
            safe_print(f"[INFO] Email {msg_id[:8]}... marcado como lido")

    # Salvar CSV
    csv_path = base_dir.parent / f"{vaga_slug}_aprovados.csv"
    with open(csv_path, "w", encoding="utf-8", newline="") as f:
        writer = csv.DictWriter(
            f, fieldnames=["arquivo", "formacoes_encontradas"]
        )
        writer.writeheader()
        for row in aprovados_info:
            writer.writerow(row)

    # Log final
    percent_aprov = (
        (total_aprovados / total_anexos * 100) if total_anexos else 0
    )
    top_formacoes = formacao_counter.most_common(5)

    safe_print(f"[RESUMO FINAL] Total lidos: {total_anexos}")
    safe_print(f"[RESUMO FINAL] Total aprovados: {total_aprovados}")
    safe_print(f"[RESUMO FINAL] % aprovação: {percent_aprov:.2f}%")
    safe_print("[RESUMO FINAL] Top 5 formações encontradas:")
    for form, count in top_formacoes:
        safe_print(f"  {form}: {count}")

    safe_print("[END] triagem ok")
    sys.exit(0)

if __name__ == "__main__":
    main()

